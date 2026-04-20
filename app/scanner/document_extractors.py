"""
Extracteurs de texte pour documents Scanner Universel — Phase 6.

Module autonome qui extrait le texte depuis :
- PDF (pdfplumber + fallback Claude Vision si pdfplumber rate)
- DOCX (python-docx)
- XLSX (openpyxl)
- Images (Claude Vision pour lire les photos avec texte,
  etiquettes, plaques signaletiques, documents pris en photo)

Retourne toujours une string de texte propre, prete a etre vectorisee
par le processor du Scanner Universel.

Usage :
    from app.scanner.document_extractors import extract_document_text

    text = extract_document_text(
        content_bytes=mes_bytes,
        filename='devis_S24068.pdf',
        mime_type='application/pdf',
    )
    # -> "Bon de commande n S24068..."

Dependances :
    pip install pdfplumber python-docx openpyxl

Variables env utilisees :
    ANTHROPIC_API_KEY (pour Claude Vision sur images)
"""

import io
import base64
import logging
from typing import Optional, Union

logger = logging.getLogger("raya.scanner.doc_extractors")

# Taille maximale de texte retourne par document (8000 caracteres =
# ~2000 tokens, limite de embed() pour OpenAI text-embedding-3-small)
MAX_TEXT_LENGTH = 8000


def extract_text_from_pdf(content_bytes: bytes,
                           max_pages: int = 100) -> Optional[str]:
    """Extrait le texte d un PDF via pdfplumber.

    Args:
        content_bytes: bytes du PDF
        max_pages: nombre max de pages a lire (evite les PDF de 500 pages
                   qui exploseraient le contexte et le cout d embedding)

    Retourne le texte concatene de toutes les pages (nettoye), ou None
    en cas d echec. Si le PDF est scanne (image pure sans couche texte),
    pdfplumber retournera une chaine vide -> le caller peut fallback sur
    Claude Vision via extract_text_from_image.
    """
    try:
        import pdfplumber
    except ImportError:
        logger.error("[DocExtract] pdfplumber non installe - "
                     "pip install pdfplumber")
        return None

    try:
        with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
            pages_to_read = min(len(pdf.pages), max_pages)
            texts = []
            for i in range(pages_to_read):
                page_text = pdf.pages[i].extract_text() or ""
                if page_text.strip():
                    texts.append(f"=== Page {i+1} ===\n{page_text.strip()}")
            full_text = "\n\n".join(texts)
            return full_text[:MAX_TEXT_LENGTH] if full_text else None
    except Exception as e:
        logger.warning("[DocExtract] PDF extraction echouee : %s", e)
        return None


def extract_text_from_docx(content_bytes: bytes) -> Optional[str]:
    """Extrait le texte d un DOCX via python-docx.

    Inclut :
    - Paragraphes du corps du document
    - Contenu des tableaux (cellule par cellule)
    - Headers/footers (si presents)

    Retourne le texte concatene ou None en cas d echec.
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.error("[DocExtract] python-docx non installe - "
                     "pip install python-docx")
        return None

    try:
        doc = DocxDocument(io.BytesIO(content_bytes))
        parts = []
        # Paragraphes
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Tableaux (contenu cellule par cellule)
        for table in doc.tables:
            for row in table.rows:
                cells_text = [cell.text.strip() for cell in row.cells
                              if cell.text.strip()]
                if cells_text:
                    parts.append(" | ".join(cells_text))
        full_text = "\n".join(parts)
        return full_text[:MAX_TEXT_LENGTH] if full_text else None
    except Exception as e:
        logger.warning("[DocExtract] DOCX extraction echouee : %s", e)
        return None


def extract_text_from_xlsx(content_bytes: bytes,
                            max_rows_per_sheet: int = 500) -> Optional[str]:
    """Extrait le texte d un XLSX via openpyxl.

    Parcourt toutes les feuilles du classeur, extrait les cellules non
    vides, retourne un texte structure :
        === Feuille : Devis ===
        A1=Client | B1=Date | C1=Montant
        A2=AZEM | B2=2024-03-15 | C2=37500

    Limite a max_rows_per_sheet par feuille (les XLSX peuvent etre tres
    gros, ex : exports compta avec 50 000 lignes).
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        logger.error("[DocExtract] openpyxl non installe - "
                     "pip install openpyxl")
        return None

    try:
        wb = load_workbook(io.BytesIO(content_bytes), data_only=True,
                           read_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Feuille : {sheet_name} ===")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                if row_count >= max_rows_per_sheet:
                    parts.append(f"... ({row_count}+ lignes, tronque)")
                    break
                non_empty = [str(c).strip() for c in row
                             if c is not None and str(c).strip()]
                if non_empty:
                    parts.append(" | ".join(non_empty))
                    row_count += 1
        full_text = "\n".join(parts)
        return full_text[:MAX_TEXT_LENGTH] if full_text else None
    except Exception as e:
        logger.warning("[DocExtract] XLSX extraction echouee : %s", e)
        return None


def extract_text_from_pdf_via_claude(content_bytes: bytes,
                                       filename: str = "") -> Optional[str]:
    """Fallback Claude pour PDF scannes (sans couche texte).

    Utilise le type `document` natif d Anthropic (pas `image`) : Claude lit
    directement le PDF, fait l OCR natif et comprend les tableaux.
    Jusqu a 100 pages / 32 Mo par PDF.
    Ref : https://docs.anthropic.com/en/docs/build-with-claude/pdf-support
    """
    import os
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        logger.warning("[DocExtract] ANTHROPIC_API_KEY absente pour PDF")
        return None
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
    except ImportError:
        logger.error("[DocExtract] anthropic SDK non installe")
        return None
    try:
        pdf_b64 = base64.standard_b64encode(content_bytes).decode("utf-8")
        response = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=4000,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "document", "source": {
                        "type": "base64",
                        "media_type": "application/pdf",
                        "data": pdf_b64,
                    }},
                    {"type": "text", "text": (
                        "Extrais TOUT le texte de ce PDF et retourne-le tel "
                        "quel, sans commentaire. Si c est un devis ou un "
                        "document structure, conserve la structure. Si "
                        "aucun texte lisible, reponds 'AUCUN_TEXTE'."
                    )},
                ],
            }],
        )
        text = response.content[0].text.strip()
        if text == "AUCUN_TEXTE" or not text:
            return None
        return text[:MAX_TEXT_LENGTH]
    except Exception as e:
        logger.warning("[DocExtract] Claude PDF echec %s : %s",
                       filename[:40], str(e)[:200])
        return None


def extract_text_from_image(content_bytes: bytes,
                             mime_type: str = "image/jpeg",
                             context_hint: str = "") -> Optional[str]:
    """Extrait le texte visible dans une image via Claude Vision.

    Utile pour :
    - Photos de plaques signaletiques (panneaux solaires, onduleurs)
    - Documents pris en photo (devis, attestations papier)
    - Etiquettes de materiel
    - Schemas ou plans annotes
    - Captures d ecran

    Args:
        content_bytes: bytes de l image
        mime_type: 'image/jpeg', 'image/png', 'image/gif', 'image/webp'
        context_hint: hint optionnel pour aider Claude
            (ex: 'Cette photo montre une plaque signaletique d onduleur')

    Retourne le texte extrait ou None en cas d echec.
    """
    import os
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        logger.warning("[DocExtract] ANTHROPIC_API_KEY absente pour Vision")
        return None
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
    except ImportError:
        logger.error("[DocExtract] anthropic SDK non installe")
        return None

    # Appel Claude Vision pour extraire le texte visible
    try:
        img_b64 = base64.standard_b64encode(content_bytes).decode("utf-8")
        prompt = (
            "Extrais TOUT le texte visible dans cette image et retourne-le "
            "tel quel, sans commentaire ni interpretation. Conserve la mise "
            "en forme approximative (lignes, colonnes). Si l image ne "
            "contient pas de texte, reponds strictement 'AUCUN_TEXTE'."
        )
        if context_hint:
            prompt = f"Contexte : {context_hint}\n\n{prompt}"
        # Utilise Claude Haiku pour rapidite + cout reduit (~3x moins cher
        # que Sonnet/Opus pour du simple OCR)
        response = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=2000,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image", "source": {
                        "type": "base64",
                        "media_type": mime_type,
                        "data": img_b64,
                    }},
                    {"type": "text", "text": prompt},
                ],
            }],
        )
        text = response.content[0].text.strip()
        if text == "AUCUN_TEXTE" or not text:
            return None
        return text[:MAX_TEXT_LENGTH]
    except Exception as e:
        logger.warning("[DocExtract] Claude Vision echec : %s", e)
        return None


# Mapping extension -> fonction d extraction
EXTENSION_TO_EXTRACTOR = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",   # tentative best-effort
    ".xlsx": "xlsx",
    ".xlsm": "xlsx",
    ".xls": "xlsx",
    ".jpg": "image", ".jpeg": "image", ".png": "image",
    ".gif": "image", ".webp": "image",
}

# Mapping mime_type -> fonction d extraction (plus fiable que l extension)
MIME_TO_EXTRACTOR = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xlsx",
    "image/jpeg": "image", "image/jpg": "image", "image/png": "image",
    "image/gif": "image", "image/webp": "image",
}


def extract_document_text(
    content_bytes: bytes,
    filename: str = "",
    mime_type: str = "",
    context_hint: str = "",
    fallback_vision_for_pdf: bool = True,
) -> Optional[str]:
    """Routeur principal : detecte le type et appelle l extracteur adequat.

    Args:
        content_bytes: contenu binaire du document
        filename: nom du fichier (utilise l extension si pas de mime)
        mime_type: mime type de preference (plus fiable que l extension)
        context_hint: hint pour Claude Vision (uniquement images)
        fallback_vision_for_pdf: si True, tente Vision page 1 du PDF
            quand pdfplumber retourne rien (cas PDF scanne image-only)

    Retourne le texte extrait ou None.
    """
    # 1. Determine le type
    extractor_type = None
    if mime_type:
        extractor_type = MIME_TO_EXTRACTOR.get(mime_type.lower())
    if not extractor_type and filename:
        ext = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""
        extractor_type = EXTENSION_TO_EXTRACTOR.get(ext)
    if not extractor_type:
        logger.info("[DocExtract] Type non supporte : filename=%s mime=%s",
                    filename, mime_type)
        return None

    # 2. Dispatch
    if extractor_type == "pdf":
        text = extract_text_from_pdf(content_bytes)
        # Fallback : PDF sans couche texte -> Claude avec type document
        if not text and fallback_vision_for_pdf:
            logger.info("[DocExtract] PDF vide, fallback Claude (type=document)")
            text = extract_text_from_pdf_via_claude(
                content_bytes, filename=filename)
        return text
    if extractor_type == "docx":
        return extract_text_from_docx(content_bytes)
    if extractor_type == "xlsx":
        return extract_text_from_xlsx(content_bytes)
    if extractor_type == "image":
        eff_mime = mime_type or "image/jpeg"
        return extract_text_from_image(content_bytes, eff_mime, context_hint)
    return None
